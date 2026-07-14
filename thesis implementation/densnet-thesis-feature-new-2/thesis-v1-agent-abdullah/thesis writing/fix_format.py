"""
Comprehensive formatting + content-insertion pass for ThesisDocument.docx.
Follows Thesis-Template-MS-1.docx conventions.
Run from the "thesis writing" directory.
"""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os, re

DOC_PATH = "ThesisDocument.docx"
FIG_DIR = "../figures"

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

doc = Document(DOC_PATH)


# ---------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------

def clear_text_children(run_elem):
    for child in list(run_elem):
        if child.tag in (W + "t", W + "br"):
            run_elem.remove(child)


def set_run_text(run, text):
    clear_text_children(run._element)
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            run._element.append(OxmlElement("w:br"))
        if part:
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = part
            run._element.append(t)


def insert_picture_at_start(p, image_path, width_in=6.0):
    new_run = p.add_run()
    new_run.add_picture(image_path, width=Inches(width_in))
    p._element.remove(new_run._element)
    pPr = p._element.find(W + "pPr")
    if pPr is not None:
        pPr.addnext(new_run._element)
    else:
        p._element.insert(0, new_run._element)
    return new_run


def build_word_table_after(after_p, rows, n_cols):
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tblPr = tbl._element.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = val
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    tbl_elem = tbl._element
    tbl_elem.getparent().remove(tbl_elem)
    after_p._element.addnext(tbl_elem)


def add_caption_paragraph_after(src_p, caption_text):
    new_p_elem = OxmlElement("w:p")
    src_p._element.addnext(new_p_elem)
    para = Paragraph(new_p_elem, src_p._parent)
    run = para.add_run(caption_text)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def set_heading(idx, level):
    doc.paragraphs[idx].style = doc.styles[f"Heading {level}"]


# ---------------------------------------------------------------
# PART A — Heading-level fixes (must run BEFORE new paragraphs are inserted)
# ---------------------------------------------------------------

# ABSTRACT body wrongly tagged Heading 1
doc.paragraphs[134].style = doc.styles["normal"]

# Chapter 1
set_heading(142, 2)

# Chapter 2
for idx in (185, 228, 234, 240, 246, 252, 259):
    set_heading(idx, 3)

# Chapter 3
set_heading(306, 2)
set_heading(318, 2)
for idx in (319, 325):
    set_heading(idx, 3)
set_heading(332, 2)
for idx in (334, 337, 341, 351):
    set_heading(idx, 3)
set_heading(357, 2)
for idx in (359, 361, 363, 366):
    set_heading(idx, 3)
set_heading(369, 2)

# Chapter 4
set_heading(377, 2)
set_heading(397, 2)
set_heading(398, 3)
for idx in (424, 428, 432, 436):
    set_heading(idx, 3)
set_heading(443, 2)
for idx in (444, 447, 451):
    set_heading(idx, 3)
for idx in (458, 465):
    set_heading(idx, 3)

# Chapter 5
set_heading(482, 2)
set_heading(491, 2)
set_heading(500, 2)
set_heading(510, 2)
set_heading(517, 3)
set_heading(524, 2)
set_heading(529, 2)
set_heading(531, 3)
set_heading(536, 3)
set_heading(549, 3)

# Chapter 6
set_heading(552, 1)
for idx in (567, 574, 581):
    set_heading(idx, 2)

print("Fixed heading levels.")


# ---------------------------------------------------------------
# PART B — Normalise chapter heading text
# ---------------------------------------------------------------

CHAPTER_TEXT = {
    136: "Chapter 1:\tIntroduction",
    180: "Chapter 2:\tLiterature Review",
    304: "Chapter 3:\tDataset",
    374: "Chapter 4:\tMethodology",
    480: "Chapter 5:\tResults",
    552: "Chapter 6:\tConclusion and Future Work",
}
for idx, new_text in CHAPTER_TEXT.items():
    p = doc.paragraphs[idx]
    runs = list(p.runs)
    if runs:
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
        set_run_text(runs[0], new_text)
    else:
        p.add_run(new_text)
print("Normalised chapter heading text.")


# ---------------------------------------------------------------
# PART C — Caption renumbering + bold + center
# ---------------------------------------------------------------

CAPTION_FIXES = [
    # Chapter 4 figure captions
    (395, "Figure 4.1: ", r"^\s*Caption:\s*"),
    (401, "Figure 4.2: ", r"^\s*Caption:\s*"),
    (408, "Figure 4.3: ", r"^\s*Caption:\s*"),
    (416, "Figure 4.4: ", r"^\s*Caption:\s*"),
    (423, "Figure 4.5: ", r"^\s*Caption:\s*"),
    (427, "Figure 4.6: ", r"^\s*Caption:\s*"),
    (435, "Figure 4.7: ", r"^\s*Caption:\s*"),
    (441, "Figure 4.8: ", r"^\s*Caption:\s*"),
    (446, "Figure 4.9: ", r"^\s*Caption:\s*"),
    (454, "Figure 4.10: ", r"^\s*Caption:\s*"),
    # Chapter 4 table captions
    (431, "Table 4.1: ", r"^\s*Caption:\s*"),
    (450, "Table 4.2: ", r"^\s*Caption:\s*"),
    # Chapter 5 figure captions
    (486, "Figure 5.1: ", r"^\s*Caption:\s*"),
    (488, "Figure 5.2: ", r"^\s*Caption:\s*"),
    (495, "Figure 5.3: ", r"^\s*Caption:\s*"),
    (497, "Figure 5.4: ", r"^\s*Caption:\s*"),
    (507, "Figure 5.5: ", r"^\s*Caption:\s*"),
    (509, "Figure 5.6: ", r"^\s*Caption:\s*"),
    (516, "Figure 5.7: ", r"^\s*Caption:\s*"),
    (522, "Figure 5.8: ", r"^\s*Caption:\s*"),
    (523, "Figure 5.9: ", r"^\s*Caption:\s*"),
    (546, "Figure 5.10: ", r"^\s*Caption:\s*"),
    (548, "Figure 5.11: ", r"^\s*Caption:\s*"),
    # Chapter 5 table captions (standardise separators)
    (484, "Table 5.1: ", r"^\s*Table\s*5\.1\s*\.\s*"),
    (493, "Table 5.2: ", r"^\s*Table\s*5\.2\s*\.\s*"),
    (504, "Table 5.3: ", r"^\s*5\.3\s*\.\s*"),
    (513, "Table 5.4: ", r"^\s*Table\s*5\.4\s*\.\s*"),
    (520, "Table 5.5: ", r"^\s*Table\s*5\.5\s*\.\s*"),
    (527, "Table 5.6: ", r"^\s*Table\s*5\.6\s*\.\s*"),
    (533, "Table 5.7: ", r"^\s*Table\s*5\.7\s*\.\s*"),
    (543, "Table 5.9: ", r"^\s*Table\s*5\.9\s*\.\s*"),
    (551, "Table 5.11: ", r"^\s*Table\s*5\.11\s*\.\s*"),
]


def rewrite_caption(idx, new_prefix, strip_regex):
    p = doc.paragraphs[idx]
    runs = list(p.runs)
    if not runs:
        return
    full_text = p.text
    leading_nl = ""
    if full_text.startswith("\n"):
        leading_nl = "\n"
        full_text = full_text.lstrip("\n")
    new_body = re.sub(strip_regex, "", full_text, count=1)
    new_text = leading_nl + new_prefix + new_body
    has_drawing_first = bool(runs[0]._element.findall(f".//{W}drawing"))
    if has_drawing_first and len(runs) > 1:
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
        new_run = p.add_run()
        set_run_text(new_run, new_text)
        new_run.bold = True
    else:
        for r in runs:
            r._element.getparent().remove(r._element)
        new_run = p.add_run()
        set_run_text(new_run, new_text)
        new_run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


for idx, prefix, strip_re in CAPTION_FIXES:
    rewrite_caption(idx, prefix, strip_re)
print(f"Rewrote {len(CAPTION_FIXES)} captions.")

# Chapter 3 captions: already say "Figure 3.X:" — just bold + center
for idx in (317, 324, 331, 350, 356):
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
print("Bolded chapter 3 figure captions.")


# ---------------------------------------------------------------
# PART D — Insert pictures into existing placeholder paragraphs
# (these don't add new paragraphs, so indices remain stable)
# ---------------------------------------------------------------

def replace_placeholder_with_image(idx, image_filename):
    p = doc.paragraphs[idx]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run()
    run.add_picture(os.path.join(FIG_DIR, image_filename), width=Inches(6.0))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


replace_placeholder_with_image(348, "augmentation_examples.png")
replace_placeholder_with_image(440, "tta_calibration.png")
replace_placeholder_with_image(453, "bootstrap_ci.png")
print("Inserted 3 missing images at placeholders.")


# ---------------------------------------------------------------
# PART E — Replace table placeholders with Word tables
# (Tables don't shift paragraph indices because they are sibling <w:tbl> elements.)
# ---------------------------------------------------------------

# Table 4.1 (hyperparameters) at para 430
p430 = doc.paragraphs[430]
for r in list(p430.runs):
    r._element.getparent().remove(r._element)
hp_rows = [
    ("Hyperparameter", "Baseline (linear head)", "v3.6 two-head"),
    ("Optimiser",            "AdamW", "AdamW"),
    ("Weight decay",         "1e-4", "5e-4"),
    ("Learning rate",        "1e-3", "1e-4 (binary) / 3e-4 (subtype)"),
    ("LR schedule",          "Cosine, 30 epochs", "3-epoch warm-up → cosine, 60 epochs"),
    ("Batch size",           "16", "16"),
    ("Loss",                 "CE + smoothing 0.1", "BCE + 1.0·LA-CE (τ=1.0, τ_Ductal=0.0)"),
    ("Class weighting",      "WeightedRandomSampler", "WeightedRandomSampler (1/√fc)"),
    ("Label smoothing",      "0.1", "0.05"),
    ("Dropout",              "—", "0.5"),
    ("EMA decay",            "—", "0.999"),
    ("Gradient clip",        "‖g‖₂ ≤ 1.0", "‖g‖₂ ≤ 1.0"),
    ("Selection criterion",  "max val bin-F1", "0.3·bin-F1 + 0.7·macro-F1, gated bin-F1>0.970"),
    ("TTA",                  "—", "10 passes, N(0, 0.01) feature noise"),
    ("Calibration",          "—", "Per-class temperature scaling (LBFGS on val NLL)"),
    ("Random seed",          "42", "42"),
]
build_word_table_after(p430, hp_rows, n_cols=3)
print("Inserted Table 4.1 (hyperparameters) after para 430.")

# Table 4.2 (mcnemar_results) at para 449
p449 = doc.paragraphs[449]
for r in list(p449.runs):
    r._element.getparent().remove(r._element)
mcn_rows = [
    ("Baseline", "Discordant b+c", "b (v3.6 wins)", "c (baseline wins)", "p-value"),
    ("Swin-B (linear head)",        "198", "154", "44", "<10⁻¹²"),
    ("ConvNeXt-B (linear head)",    "175", "136", "39", "<10⁻⁹"),
    ("Logit ensemble",              "160", "124", "36", "<10⁻¹⁰"),
    ("Feature ensemble (MLP)",      "130", "101", "29", "<10⁻⁷"),
    ("Binary-opt fusion (Var. A)",   "75",  "58", "17", "<10⁻⁵"),
]
build_word_table_after(p449, mcn_rows, n_cols=5)
print("Inserted Table 4.2 (mcnemar) after para 449.")


# ---------------------------------------------------------------
# PART F — Insert Table 3.2 (gating ablation) — adds NEW paragraph (caption)
# Run LAST so paragraph-index work above is unaffected.
# ---------------------------------------------------------------
p418 = doc.paragraphs[418]
cap_3_2 = add_caption_paragraph_after(
    p418,
    "Table 3.2: Gating-mechanism ablation under patient-level 5-fold CV. "
    "Macro-F1 (mean ± std) and parameter counts."
)
gating_rows = [
    ("Variant", "Macro-F1 mean", "Std", "Params", "Gate weights"),
    ("No gate (feature ensemble baseline)",         "0.8211", "0.1599", "562 K", "—"),
    ("Hard gate (argmax routing on v3.6 experts)",  "0.7959", "0.0932", "4.27 M", "∈{0,1}"),
    ("Soft average (w_bin = w_8c = 0.5)",           "0.8129", "0.0923", "4.27 M", "0.5"),
    ("Two experts, no gate (logit avg.)",           "0.8054", "0.1142", "4.13 M", "0.5"),
    ("Learned gate (v3.6, final)",                  "0.8352", "0.0876", "4.27 M", "data-dep."),
]
build_word_table_after(cap_3_2, gating_rows, n_cols=5)
print("Inserted Table 3.2 (gating ablation) after caption following para 418.")


# ---------------------------------------------------------------
# Save
# ---------------------------------------------------------------
doc.save(DOC_PATH)
print(f"\nSaved {DOC_PATH}.")
