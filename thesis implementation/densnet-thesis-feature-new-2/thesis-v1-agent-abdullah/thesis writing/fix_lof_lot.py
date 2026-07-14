"""
Rebuild LIST OF TABLES and LIST OF FIGURES with all current captions
(including the 5 tables just inserted) and cleaner format.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

DOC_PATH = "ThesisDocument.docx"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

doc = Document(DOC_PATH)


def find_para(needle):
    for i, p in enumerate(doc.paragraphs):
        if needle in p.text:
            return i, p
    return -1, None


def collect_captions(prefix_regex):
    """Find caption paragraphs ordered by document position."""
    out = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.lstrip("\n").strip()
        if re.match(prefix_regex, text):
            out.append((i, text))
    return out


def remove_section_entries(start_idx, end_idx):
    """Remove paragraphs in (start_idx, end_idx) keeping 'Page No.' on start_idx+1."""
    # Re-resolve indices each iteration because deletions shift positions.
    removed = 0
    while True:
        end_para = doc.paragraphs[end_idx] if end_idx < len(doc.paragraphs) else None
        if end_para is None:
            break
        # Find first paragraph between start (after Page No.) and end_para that exists
        first_idx = start_idx + 2  # skip header + "Page No."
        if first_idx >= end_idx:
            break
        # Remove that paragraph
        p = doc.paragraphs[first_idx]
        p._element.getparent().remove(p._element)
        end_idx -= 1
        removed += 1
    return removed


def build_lof_entry(label, body):
    """Build a paragraph with bold label, plain body, right-aligned dot leader, '—' placeholder."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9000")
    tabs.append(tab)
    pPr.append(tabs)
    # 0.25" hanging indent so wrapped lines align under text
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "720")
    pPr.append(ind)
    p.append(pPr)

    # Run 1: bold label "Figure X.Y:"
    r1 = OxmlElement("w:r")
    rPr1 = OxmlElement("w:rPr"); b = OxmlElement("w:b"); rPr1.append(b); r1.append(rPr1)
    t1 = OxmlElement("w:t"); t1.set(qn("xml:space"), "preserve"); t1.text = label; r1.append(t1)
    p.append(r1)

    # Run 2: plain body + tab + page placeholder
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t"); t2.set(qn("xml:space"), "preserve"); t2.text = " " + body
    r2.append(t2)
    tab_r = OxmlElement("w:tab")
    r2.append(tab_r)
    t3 = OxmlElement("w:t"); t3.set(qn("xml:space"), "preserve"); t3.text = "—"
    r2.append(t3)
    p.append(r2)

    return p


def split_caption(text):
    """'Figure 4.1: Description...' -> ('Figure 4.1:', 'Description...')"""
    m = re.match(r"^((?:Figure|Table)\s+\d+(?:\.\d+)?(?:[a-zA-Z])?:)\s*(.*)$", text)
    if m:
        return m.group(1), m.group(2).strip()
    return text, ""


# ---------------------------------------------------------------
# Locate section anchors
# ---------------------------------------------------------------
i_lot, _ = find_para("LIST OF TABLES")
i_lof, _ = find_para("LIST OF FIGURES")
i_los, _ = find_para("LIST OF SYMBOLS, ABBREVIATIONS")
assert i_lot != -1 and i_lof != -1 and i_los != -1, "Couldn't locate LOT/LOF/LOS anchors"

# Clear current LOT entries (between i_lot+2 and i_lof)
removed_lot = remove_section_entries(i_lot, i_lof)
print(f"Cleared {removed_lot} stale LOT entries.")

# After deletion, re-resolve LOF and LOS
i_lof, _ = find_para("LIST OF FIGURES")
i_los, _ = find_para("LIST OF SYMBOLS, ABBREVIATIONS")
removed_lof = remove_section_entries(i_lof, i_los)
print(f"Cleared {removed_lof} stale LOF entries.")

# Now collect captions from the doc
tables_list = collect_captions(r"^Table\s+\d+\.\d+[a-zA-Z]?\s*:")
figures_list = collect_captions(r"^Figure\s+\d+\.\d+[a-zA-Z]?\s*:")
print(f"Collected {len(tables_list)} tables and {len(figures_list)} figures.")


# ---------------------------------------------------------------
# Insert LOT entries
# ---------------------------------------------------------------
i_lot, p_lot = find_para("LIST OF TABLES")
page_no_p = doc.paragraphs[i_lot + 1]  # "Page No." paragraph
anchor = page_no_p._element
for _, text in tables_list:
    label, body = split_caption(text)
    entry = build_lof_entry(label, body)
    anchor.addnext(entry)
    anchor = entry
print(f"Inserted {len(tables_list)} LOT entries.")

# ---------------------------------------------------------------
# Insert LOF entries
# ---------------------------------------------------------------
i_lof, p_lof = find_para("LIST OF FIGURES")
page_no_p = doc.paragraphs[i_lof + 1]  # "Page No." paragraph
anchor = page_no_p._element
for _, text in figures_list:
    label, body = split_caption(text)
    entry = build_lof_entry(label, body)
    anchor.addnext(entry)
    anchor = entry
print(f"Inserted {len(figures_list)} LOF entries.")


doc.save(DOC_PATH)
print(f"\nSaved {DOC_PATH}.")
