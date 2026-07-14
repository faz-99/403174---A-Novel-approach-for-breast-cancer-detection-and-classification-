"""
TOC / LOT / LOF + Full Deletion-AUC table insertion.
Run from "thesis writing" directory.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
import re

DOC_PATH = "ThesisDocument.docx"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

doc = Document(DOC_PATH)


# ---------------------------------------------------------------
# Build a Word TOC field paragraph
# ---------------------------------------------------------------

def build_toc_paragraph(toc_field_code, placeholder_text):
    """Return a <w:p> element containing a complex field (TOC ...).
    Word will populate it when the user opens the file (because dirty=true).
    """
    p = OxmlElement("w:p")

    def add_run(parent, *children):
        r = OxmlElement("w:r")
        for c in children:
            r.append(c)
        parent.append(r)
        return r

    # Field begin (dirty=true so Word recomputes on open)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    add_run(p, fld_begin)

    # Field instruction text
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " " + toc_field_code + " "
    add_run(p, instr)

    # Field separator
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    add_run(p, fld_sep)

    # Placeholder visible until Word refreshes
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = placeholder_text
    add_run(p, t)

    # Field end
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    add_run(p, fld_end)

    return p


def insert_after(src_p, *elems):
    cur = src_p._element
    for e in elems:
        cur.addnext(e)
        cur = e


# ---------------------------------------------------------------
# 1) TABLE OF CONTENTS — replace empty paragraphs 0098 with TOC field
# ---------------------------------------------------------------

toc_heading_p = doc.paragraphs[97]   # "TABLE OF CONTENTS"
# Remove the existing blank paragraph at 98 and replace with TOC field
blank_98 = doc.paragraphs[98]
blank_98._element.getparent().remove(blank_98._element)
# Build TOC field paragraph
toc_p = build_toc_paragraph(
    'TOC \\o "1-3" \\h \\z \\u',
    "Right-click here and choose Update Field (or press F9) to populate the Table of Contents."
)
toc_heading_p._element.addnext(toc_p)
print("Inserted TOC field after 'TABLE OF CONTENTS'.")

# Refresh paragraph indices after removing/inserting
doc_paras = doc.paragraphs


def find_paragraph_by_text(needle, equals=False):
    for i, p in enumerate(doc.paragraphs):
        if equals:
            if p.text.strip() == needle:
                return i, p
        else:
            if needle in p.text:
                return i, p
    return -1, None


# ---------------------------------------------------------------
# 2) LIST OF TABLES — build static list from existing "Table X.Y:" captions
# ---------------------------------------------------------------

def collect_captions(prefix_regex):
    """Return list of (paragraph_idx, full_caption_text) ordered by appearance."""
    out = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.lstrip("\n").strip()
        if re.match(prefix_regex, text):
            out.append((i, text))
    return out


tables_list = collect_captions(r"^Table\s+\d+\.\d+\s*:")
figures_list = collect_captions(r"^Figure\s+\d+\.\d+\s*:")

print(f"Found {len(tables_list)} Table captions and {len(figures_list)} Figure captions.")


def build_listing_paragraph(text, style="normal"):
    """Create a <w:p> with a single run of plain text in given style."""
    p = OxmlElement("w:p")
    # pPr - apply tab leader (right-aligned tab at 8550 with dotted leader)
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style)
    pPr.append(pStyle)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "8550")
    tabs.append(tab)
    pPr.append(tabs)
    p.append(pPr)
    # Run
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def truncate(s, n=110):
    s = s.replace("\n", " ")
    if len(s) > n:
        return s[:n].rstrip() + "…"
    return s


# Insert LIST OF TABLES content after "Page No." paragraph (idx 100 originally — but
# may have shifted due to TOC blank removal; locate by content)
idx_lot_header, _ = find_paragraph_by_text("LIST OF TABLES", equals=False)
idx_lot_page_no = idx_lot_header + 1  # "Page No."
lot_page_p = doc.paragraphs[idx_lot_page_no]

# Remove any empty paragraphs between LOT heading and LIST OF FIGURES heading
idx_lof_header, _ = find_paragraph_by_text("LIST OF FIGURES", equals=False)
# Wipe empty paragraphs between idx_lot_page_no+1 and idx_lof_header (exclusive)
removed = 0
i = idx_lot_page_no + 1
while i < idx_lof_header:
    p = doc.paragraphs[i]
    if p.text.strip() == "":
        p._element.getparent().remove(p._element)
        idx_lof_header -= 1
        removed += 1
    else:
        i += 1
print(f"Removed {removed} empty paras between LOT/LOF headers.")

# Insert table entries
prev_elem = doc.paragraphs[idx_lot_page_no]._element
for _, text in tables_list:
    entry = build_listing_paragraph(truncate(text))
    prev_elem.addnext(entry)
    prev_elem = entry
print(f"Inserted {len(tables_list)} entries under LIST OF TABLES.")


# ---------------------------------------------------------------
# 3) LIST OF FIGURES
# ---------------------------------------------------------------

idx_lof_header, _ = find_paragraph_by_text("LIST OF FIGURES", equals=False)
idx_lof_page_no = idx_lof_header + 1  # "Page No."

# Find next major heading or non-empty paragraph to bound the LOF section
# LOF section ends just before "LIST OF SYMBOLS, ABBREVIATIONS and acronyms"
idx_los_header, _ = find_paragraph_by_text("LIST OF SYMBOLS, ABBREVIATIONS", equals=False)

removed = 0
i = idx_lof_page_no + 1
while i < idx_los_header:
    p = doc.paragraphs[i]
    if p.text.strip() == "":
        p._element.getparent().remove(p._element)
        idx_los_header -= 1
        removed += 1
    else:
        i += 1
print(f"Removed {removed} empty paras between LOF and LOS headers.")

prev_elem = doc.paragraphs[idx_lof_page_no]._element
for _, text in figures_list:
    entry = build_listing_paragraph(truncate(text))
    prev_elem.addnext(entry)
    prev_elem = entry
print(f"Inserted {len(figures_list)} entries under LIST OF FIGURES.")


# ---------------------------------------------------------------
# 4) Full deletion-AUC table — para that references results/table_4_2.json
#    Insert AFTER the paragraph mentioning it (text contains 'table_4_2.json').
# ---------------------------------------------------------------

def build_word_table_after(after_elem, rows, n_cols):
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tblPr = tbl._element.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = val
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    tbl_elem = tbl._element
    tbl_elem.getparent().remove(tbl_elem)
    after_elem.addnext(tbl_elem)


idx_519, _ = find_paragraph_by_text("table_4_2.json", equals=False)
if idx_519 != -1:
    # Insert a caption + full table immediately after that paragraph
    src_p = doc.paragraphs[idx_519]
    cap_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); pPr.append(jc)
    cap_p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b"); rPr.append(b)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = "Table 5.4a: Full deletion-AUC across all five attribution methods and three models, with Fusion std, runtime, and paired p-value vs. Integrated Gradients on the Fusion model."
    r.append(t)
    cap_p.append(r)
    src_p._element.addnext(cap_p)

    full_rows = [
        ("Method",              "Swin",  "ConvNeXt", "Fusion", "Fusion std", "Runtime (s)", "vs IG (p)"),
        ("Grad-CAM",            "0.502", "0.382",    "0.482",  "0.082",      "0.052",       "0.020"),
        ("Grad-CAM++",          "0.488", "0.442",    "0.462",  "0.081",      "0.054",       "0.025"),
        ("HiResCAM",            "0.535", "0.382",    "0.482",  "0.083",      "0.052",       "0.020"),
        ("LayerCAM",            "0.525", "0.443",    "0.530",  "0.087",      "0.051",       "0.012"),
        ("Integrated Gradients","0.409", "0.130",    "0.178",  "0.042",      "0.330",       "—"),
    ]
    build_word_table_after(cap_p, full_rows, n_cols=7)
    print("Inserted full deletion-AUC table after para mentioning table_4_2.json.")
else:
    print("Could not locate table_4_2.json reference.")


# ---------------------------------------------------------------
# Save
# ---------------------------------------------------------------
doc.save(DOC_PATH)
print(f"\nSaved {DOC_PATH}.")
