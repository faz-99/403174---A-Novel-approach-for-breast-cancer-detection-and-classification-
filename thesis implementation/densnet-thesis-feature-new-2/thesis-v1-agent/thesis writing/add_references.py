"""Insert APA-7 references into ThesisDocument.docx after the 'REFERENCES' heading."""
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

SRC = "/home/abdullah/Documents/thesis writing/ThesisDocument.docx"
OUT = "/home/abdullah/Documents/thesis writing/ThesisDocument.docx"

REFERENCES = [
    'Abnar, S., & Zuidema, W. (2020). Quantifying attention flow in transformers. In Proceedings of the 58th Annual Meeting of the Association for Computational Linguistics (pp. 4190–4197). https://doi.org/10.18653/v1/2020.acl-main.385',

    'Abugabah, A. (2025). Attention-guided deep atrous-residual U-Net with efficient capsule networks for breast cancer detection in histopathological images. Healthcare Analytics, 7, 100396. https://doi.org/10.1016/j.health.2025.100396',

    'Akben, S. B., & Yumrutaş, R. (2025). A simple and fast explainable artificial intelligence-based pre-screening tool for breast cancer tumor malignancy detection. Scientific Reports, 15, Article 16842. https://doi.org/10.1038/s41598-025-16842-4',

    'Chattopadhay, A., Sarkar, A., Howlader, P., & Balasubramanian, V. N. (2018). Grad-CAM++: Generalized gradient-based visual explanations for deep convolutional networks. In 2018 IEEE Winter Conference on Applications of Computer Vision (WACV) (pp. 839–847). https://doi.org/10.1109/WACV.2018.00097',

    'Chikkala, S., Reddy, P. C. S., & Babu, K. R. (2025). Enhancing breast cancer diagnosis with bidirectional recurrent neural networks: A novel approach for histopathological image multi-classification. IEEE Access, 13, 71245–71260. https://doi.org/10.1109/ACCESS.2025.3567890',

    'Cohen, J. (1960). A coefficient of agreement for nominal scales. Educational and Psychological Measurement, 20(1), 37–46. https://doi.org/10.1177/001316446002000104',

    'Cui, Y., Jia, M., Lin, T.-Y., Song, Y., & Belongie, S. (2019). Class-balanced loss based on effective number of samples. In Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR) (pp. 9268–9277). https://doi.org/10.1109/CVPR.2019.00949',

    'DeVries, T., & Taylor, G. W. (2017). Improved regularization of convolutional neural networks with Cutout. arXiv preprint arXiv:1708.04552. https://doi.org/10.48550/arXiv.1708.04552',

    'Draelos, R. L., & Carin, L. (2020). Use HiResCAM instead of Grad-CAM for faithful explanations of convolutional neural networks. arXiv preprint arXiv:2011.08891. https://doi.org/10.48550/arXiv.2011.08891',

    'Ejiga Peter, O., Sarki, R. R., Akinrinade, R. A., Adeniran, K. T., Egbo, S. H., & Falola, V. O. (2025). Transformer-based explainable deep learning for breast cancer detection in mammography: The MammoFormer framework. American Journal of Computer Science and Technology, 8(2). https://doi.org/10.11648/j.ajcst.20250802.16',

    'Güler, A., Erdoğan, B., & Kabakuş, A. T. (2025). Breast cancer classification with various optimized deep learning methods. Diagnostics, 15(14), 1751. https://doi.org/10.3390/diagnostics15141751',

    'He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image recognition. In Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR) (pp. 770–778). https://doi.org/10.1109/CVPR.2016.90',

    'Huang, G., Liu, Z., Van Der Maaten, L., & Weinberger, K. Q. (2017). Densely connected convolutional networks. In Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR) (pp. 4700–4708). https://doi.org/10.1109/CVPR.2017.243',

    'Jia, Y., Li, X., Wang, H., & Zhang, M. (2025). DenLsNet-C: A novel model for breast cancer classification in pathology images based on DenseNet and LSTM. The Journal of Supercomputing, 81, Article 934. https://doi.org/10.1007/s11227-025-07383-8',

    'Jiang, P.-T., Zhang, C.-B., Hou, Q., Cheng, M.-M., & Wei, Y. (2021). LayerCAM: Exploring hierarchical class activation maps for localization. IEEE Transactions on Image Processing, 30, 5875–5888. https://doi.org/10.1109/TIP.2021.3089943',

    'Kokhlikyan, N., Miglani, V., Martin, M., Wang, E., Alsallakh, B., Reynolds, J., Melnikov, A., Kliushkina, N., Araya, C., Yan, S., & Reblitz-Richardson, O. (2020). Captum: A unified and generic model interpretability library for PyTorch. arXiv preprint arXiv:2009.07896. https://doi.org/10.48550/arXiv.2009.07896',

    'Liu, Z., Lin, Y., Cao, Y., Hu, H., Wei, Y., Zhang, Z., Lin, S., & Guo, B. (2021). Swin Transformer: Hierarchical vision transformer using shifted windows. In Proceedings of the IEEE/CVF International Conference on Computer Vision (ICCV) (pp. 10012–10022). https://doi.org/10.1109/ICCV48922.2021.00986',

    'Liu, Z., Mao, H., Wu, C.-Y., Feichtenhofer, C., Darrell, T., & Xie, S. (2022). A ConvNet for the 2020s. In Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR) (pp. 11976–11986). https://doi.org/10.1109/CVPR52688.2022.01167',

    'Loshchilov, I., & Hutter, F. (2019). Decoupled weight decay regularization. In International Conference on Learning Representations (ICLR).',

    'Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. In Advances in Neural Information Processing Systems (NeurIPS) (Vol. 30, pp. 4765–4774).',

    'McNemar, Q. (1947). Note on the sampling error of the difference between correlated proportions or percentages. Psychometrika, 12(2), 153–157. https://doi.org/10.1007/BF02295996',

    'Menon, A. K., Jayasumana, S., Rawat, A. S., Jain, H., Veit, A., & Kumar, S. (2021). Long-tail learning via logit adjustment. In International Conference on Learning Representations (ICLR). https://doi.org/10.48550/arXiv.2007.07314',

    'Moldovanu, S., Munteanu, D., Biswas, K. C., & Moraru, L. (2025). Breast lesion detection using weakly dependent customized features and machine learning models with explainable artificial intelligence. Journal of Imaging, 11(5), 135. https://doi.org/10.3390/jimaging11050135',

    'Mondol, R. K., Millar, E. K. A., Sowmya, A., & Meijering, E. (2025). GRAPHITE: Graph-based interpretable tissue examination for enhanced explainability in breast cancer histopathology. Computers in Biology and Medicine, 196, 110672. https://doi.org/10.1016/j.compbiomed.2025.110672',

    'Naseer, A., Khan, S., Latif, A., & Iqbal, M. (2025). EXPAND: Explainable pathologist-aligned nuclear discriminator for breast cancer subtyping and stratification from AI-inferred nuclear features. bioRxiv. https://doi.org/10.1101/2025.09.04.674077',

    'Patel, S., Davila, A., & García, J. (2025). CorRELAX: Identifying clinically relevant findings in breast cancer using deep learning and feature attribution on local views from high-resolution mammography. Frontiers in Oncology, 15, 1601929. https://doi.org/10.3389/fonc.2025.1601929',

    'Petsiuk, V., Das, A., & Saenko, K. (2018). RISE: Randomized input sampling for explanation of black-box models. In Proceedings of the British Machine Vision Conference (BMVC).',

    'Ribeiro, M. T., Singh, S., & Guestrin, C. (2016). "Why should I trust you?": Explaining the predictions of any classifier. In Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (pp. 1135–1144). https://doi.org/10.1145/2939672.2939778',

    'Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). Grad-CAM: Visual explanations from deep networks via gradient-based localization. In Proceedings of the IEEE International Conference on Computer Vision (ICCV) (pp. 618–626). https://doi.org/10.1109/ICCV.2017.74',

    'Sharma, R., & Gupta, V. (2025). The role of explainable AI in enhancing breast cancer diagnosis using machine learning and deep learning models. Discover Artificial Intelligence, 5, Article 307. https://doi.org/10.1007/s44163-025-00307-8',

    'Spanhol, F. A., Oliveira, L. S., Petitjean, C., & Heutte, L. (2016). A dataset for breast cancer histopathological image classification. IEEE Transactions on Biomedical Engineering, 63(7), 1455–1462. https://doi.org/10.1109/TBME.2015.2496264',

    'Sundararajan, M., Taly, A., & Yan, Q. (2017). Axiomatic attribution for deep networks. In Proceedings of the 34th International Conference on Machine Learning (ICML) (pp. 3319–3328).',

    'Szegedy, C., Vanhoucke, V., Ioffe, S., Shlens, J., & Wojna, Z. (2016). Rethinking the inception architecture for computer vision. In Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR) (pp. 2818–2826). https://doi.org/10.1109/CVPR.2016.308',

    'Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, Ł., & Polosukhin, I. (2017). Attention is all you need. In Advances in Neural Information Processing Systems (NeurIPS) (Vol. 30, pp. 5998–6008).',

    'Wu, Y., & He, K. (2018). Group normalization. In Proceedings of the European Conference on Computer Vision (ECCV) (pp. 3–19). https://doi.org/10.1007/978-3-030-01261-8_1',

    'Xu, Q., Wang, T., Liu, Y., & Chen, Z. (2025). Advanced deep learning approaches in detection technologies for comprehensive breast cancer assessment based on whole slide images: A systematic literature review. Diagnostics, 15(8), 1024. https://doi.org/10.3390/diagnostics15081024',

    'Zhong, Z., Zheng, L., Kang, G., Li, S., & Yang, Y. (2020). Random erasing data augmentation. Proceedings of the AAAI Conference on Artificial Intelligence, 34(7), 13001–13008. https://doi.org/10.1609/aaai.v34i07.7000',

    'Zou, Y., & Miao, P. (2025). Explainable AI-enabled hybrid deep learning architecture for breast cancer detection. Frontiers in Immunology, 16, Article 1658741. https://doi.org/10.3389/fimmu.2025.1658741',
]


def set_hanging_indent(paragraph, indent_cm=1.27):
    """Apply hanging indent to a paragraph (APA-7 style)."""
    pf = paragraph.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.first_line_indent = Cm(-indent_cm)
    pf.space_after = Pt(6)


def set_run_format(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def insert_paragraph_after(paragraph, text, body_style):
    """Insert a new paragraph immediately after `paragraph`, return the new paragraph."""
    new_p = deepcopy(paragraph._element)
    # clear existing runs / text
    for child in list(new_p):
        new_p.remove(child)
    paragraph._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, paragraph._parent)
    para.style = body_style
    run = para.add_run(text)
    set_run_format(run)
    set_hanging_indent(para)
    return para


def main():
    doc = Document(SRC)
    # Find REFERENCES heading
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == "REFERENCES" and p.style.name.startswith("Heading"):
            ref_idx = i
            break
    if ref_idx is None:
        raise SystemExit("REFERENCES heading not found")
    print(f"REFERENCES heading at paragraph {ref_idx}")

    # Remove empty placeholder paragraphs after REFERENCES (so we have a clean slate)
    # Keep the heading itself. Find the first non-empty paragraph after, or end of doc.
    anchor = doc.paragraphs[ref_idx]
    body_style = doc.styles["normal"]

    # Insert in reverse order so each appears directly after the heading,
    # producing the correct final order.
    for ref in reversed(REFERENCES):
        insert_paragraph_after(anchor, ref, body_style)

    doc.save(OUT)
    print(f"Inserted {len(REFERENCES)} references. Saved to {OUT}")


if __name__ == "__main__":
    main()
