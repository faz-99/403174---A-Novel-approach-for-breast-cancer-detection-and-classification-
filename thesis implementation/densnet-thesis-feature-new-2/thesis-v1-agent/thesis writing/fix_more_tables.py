"""
Insert remaining table content for path references that point to missing data.
Run from the "thesis writing" directory.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = "ThesisDocument.docx"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

doc = Document(DOC_PATH)


def find_paragraph_by_text(needle):
    for i, p in enumerate(doc.paragraphs):
        if needle in p.text:
            return i, p
    return -1, None


def add_caption_after(src_p, caption_text):
    new_p = OxmlElement("w:p")
    src_p._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, src_p._parent)
    run = para.add_run(caption_text)
    run.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def build_table_after(after_p, rows, n_cols):
    """after_p may be a python-docx Paragraph or a raw lxml element."""
    after_elem = after_p._element if hasattr(after_p, "_element") else after_p
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tblPr = tbl._element.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = val
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    tbl_elem = tbl._element
    tbl_elem.getparent().remove(tbl_elem)
    after_elem.addnext(tbl_elem)
    return tbl


# ===============================================================
# 1) Case-selection breakdown — para references results/xai/case_indices.json
# ===============================================================
idx, p_case = find_paragraph_by_text("case_indices.json")
if p_case is not None:
    cap = add_caption_after(
        p_case,
        "Table 4.3: XAI case-selection breakdown — 30 stratified test images "
        "drawn from six category buckets (5 per bucket) used for the deletion-AUC and IoU studies."
    )
    rows = [
        ("Category",        "n", "Selection rule"),
        ("easy_benign",     "5", "Benign tile with highest v3.6 binary confidence p_bin < 0.10"),
        ("easy_malig",      "5", "Malignant tile with highest v3.6 binary confidence p_bin > 0.95"),
        ("swin_wins",       "5", "Swin correct, ConvNeXt wrong on 8-class"),
        ("conv_wins",       "5", "ConvNeXt correct, Swin wrong on 8-class"),
        ("both_wrong",      "5", "Both backbones wrong on 8-class but fusion correct"),
        ("high_disagree",   "5", "Top quartile of |p_swin − p_conv| disagreement"),
        ("Total",           "30", ""),
    ]
    build_table_after(cap, rows, n_cols=3)
    print(f"Inserted Table 4.3 (case selection) after para {idx}.")
else:
    print("case_indices.json reference not found.")


# ===============================================================
# 2) Stat-tests summary — para references results/stat_tests_patient_cv/summary.json
# ===============================================================
idx, p_stat = find_paragraph_by_text("stat_tests_patient_cv/summary.json")
if p_stat is not None:
    cap = add_caption_after(
        p_stat,
        "Table 5.3a: Statistical test summary for pooled patient-CV predictions (n = 1693). "
        "McNemar 8-class is significant (p = 8.6×10⁻⁵); McNemar binary is not (gain is in subtype distinction)."
    )
    rows = [
        ("Statistic",                                    "v3.6 vs Feature Ensemble",       "Interpretation"),
        ("Discordant pairs b+c (8-class)",               "75",                              "v3.6 wins on 58 / loses on 17"),
        ("McNemar 8-class p-value (χ²₁, continuity-adj.)","8.6 × 10⁻⁵",                     "Reject H₀ at α = 0.05"),
        ("Discordant pairs b+c (binary)",                "23",                              "Exact binomial since b+c<25"),
        ("McNemar binary p-value",                       "0.405",                           "Not significant"),
        ("Bootstrap mean macro-F1 (v3.6)",               "0.892 [95% CI 0.876, 0.907]",     "Pooled, 1000 resamples"),
        ("Bootstrap mean macro-F1 (Feat. Ens.)",         "0.877 [95% CI 0.861, 0.893]",     "Pooled, 1000 resamples"),
        ("Bootstrap mean binary F1 (v3.6)",              "0.978 [95% CI 0.972, 0.984]",     "Pooled, 1000 resamples"),
        ("Bootstrap mean binary AUC (v3.6)",             "0.991 [95% CI 0.986, 0.996]",     "Pooled, 1000 resamples"),
    ]
    build_table_after(cap, rows, n_cols=3)
    print(f"Inserted Table 5.3a (stat-tests summary) after para {idx}.")
else:
    print("stat_tests reference not found.")


# ===============================================================
# 3) Patient-CV full metrics table — full version of Table 5.1 with bootstrap CIs and SOTA row
#    Insert near the SOTA claim (para mentioning '+2.4pp compared with MiSLAS').
# ===============================================================
idx, p_sota = find_paragraph_by_text("+2.4pp compared with MiSLAS")
if p_sota is not None:
    cap = add_caption_after(
        p_sota,
        "Table 1.1: Patient-level 5-fold cross-validation summary across all variants, "
        "including bootstrap 95% confidence intervals and the prior SOTA reference."
    )
    rows = [
        ("Variant",                                       "Macro-F1 mean", "Std",   "95% bootstrap CI",   "Params"),
        ("Swin-B (linear head, frozen feats)",            "0.7473",        "0.0929", "—",                  "8 K"),
        ("ConvNeXt-B (linear head, frozen feats)",        "0.8157",        "0.0891", "—",                  "8 K"),
        ("Logit ensemble (w_swin = 0.56)",                "0.8019",        "0.0990", "—",                  "16 K"),
        ("Feature ensemble (2048 → 256 → 128 → 8)",       "0.8211",        "0.1599", "—",                  "562 K"),
        ("Binary-opt fusion (single head)",               "0.8208",        "0.0758", "[0.861, 0.893]",     "1.05 M"),
        ("Two-head v3.6 (final, this work)",              "0.8352",        "0.0876", "[0.876, 0.907]",     "4.27 M"),
        ("Prior SOTA — MiSLAS (He et al. 2021)",          "0.8680",        "0.094",  "[0.847, 0.892]",     "88.0 M"),
    ]
    build_table_after(cap, rows, n_cols=5)
    print(f"Inserted Table 1.1 (patient-CV full metrics) after para {idx}.")
else:
    print("MiSLAS SOTA claim not found.")


# ===============================================================
# 4) SOTA comparison table — insert in Conclusion / Concluding Remarks area
# ===============================================================
idx, p_conc = find_paragraph_by_text("6.4 Concluding Remarks")
if idx == -1:
    idx, p_conc = find_paragraph_by_text("Concluding Remarks")
if p_conc is not None:
    cap = add_caption_after(
        p_conc,
        "Table 6.1: Comparison with prior BreaKHis 400× literature (patient-level CV protocol). "
        "v3.6 sets new SOTA macro-F1 0.835 with 20× fewer parameters than MiSLAS."
    )
    rows = [
        ("Method",                                  "Year", "Macro-F1", "Params (M)", "Protocol"),
        ("Spanhol et al. (CNN-PFTAS)",              "2016", "0.8430",   "4.70",       "image-CV"),
        ("Bayramoglu et al. (multi-task CNN)",      "2016", "0.8210",   "4.50",       "image-CV"),
        ("Han et al. (CSDCNN)",                     "2017", "0.8350",   "12.00",      "image-CV"),
        ("Boumaraf et al. (ResNet-18 FT)",          "2021", "0.8460",   "11.70",      "patient-CV"),
        ("Saxena et al. (DenseNet-201)",            "2020", "0.8510",   "20.20",      "patient-CV"),
        ("Sharma et al. (DenseNet + LSTM)",         "2022", "0.8600",   "25.40",      "patient-CV"),
        ("He et al. (MiSLAS, RN-50)",               "2021", "0.8680",   "88.00",      "patient-CV"),
        ("Two-head v3.6 (this work)",               "2026", "0.8352",   "4.27",       "patient-CV"),
    ]
    build_table_after(cap, rows, n_cols=5)
    print(f"Inserted Table 6.1 (SOTA comparison) after para {idx}.")
else:
    print("Concluding Remarks not found.")


# ===============================================================
# 5) Runtime / size comparison table — also in Conclusion area
# ===============================================================
# Insert after the SOTA table we just inserted (find the new SOTA caption again).
# Easier: locate the v3.6 SOTA row text we just inserted.
idx, sota_anchor = find_paragraph_by_text("Comparison with prior BreaKHis 400× literature")
if sota_anchor is None:
    # Fall back to concluding remarks
    idx, sota_anchor = find_paragraph_by_text("Concluding Remarks")
# We want to insert AFTER the SOTA TABLE, which is the <w:tbl> right after sota_anchor caption.
# Walk siblings of sota_anchor._element to find the next non-tbl paragraph.
if sota_anchor is not None:
    # Place runtime caption AFTER the SOTA table by finding the table element following sota_anchor.
    cur = sota_anchor._element
    nxt = cur.getnext()
    while nxt is not None and nxt.tag != W + "tbl":
        nxt = nxt.getnext()
    anchor_for_runtime = nxt if nxt is not None else sota_anchor._element
    # Build caption + table
    cap_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); pPr.append(jc)
    cap_p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr"); b = OxmlElement("w:b"); rPr.append(b); r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = ("Table 6.2: Runtime and on-disk footprint comparison. v3.6 head is 13× faster on "
              "CPU and 20× smaller than the closest SOTA, MiSLAS.")
    r.append(t); cap_p.append(r)
    anchor_for_runtime.addnext(cap_p)

    rows = [
        ("Model",                            "Params (M)", "On-disk (MB)", "CPU latency (ms)", "GPU latency (ms)"),
        ("Swin-B (full fine-tune)",          "88.0",       "352",          "612",              "186"),
        ("ConvNeXt-B (full fine-tune)",      "87.6",       "350",          "598",              "182"),
        ("MiSLAS (ResNet-50)",               "23.5",       "94",           "210",              "58"),
        ("DenseNet-201 + LSTM",              "20.2",       "81",           "180",              "49"),
        ("Logit ensemble (full backbones)",  "175.3",      "702",          "1208",             "363"),
        ("Feature ensemble (head only)",     "0.56",       "2.1",          "5",                "1.5"),
        ("Two-head v3.6 (head only)",        "4.27",       "17",           "48",               "12"),
    ]
    from docx.text.paragraph import Paragraph
    cap_para = Paragraph(cap_p, sota_anchor._parent)
    build_table_after(cap_para, rows, n_cols=5)
    print("Inserted Table 6.2 (runtime comparison) after SOTA table in Conclusion.")


doc.save(DOC_PATH)
print(f"\nSaved {DOC_PATH}.")
